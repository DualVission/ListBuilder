from docx import Document
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_LEADER, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shape import InlineShape
import docx

from datetime import date
import os

from data.itemTypes import *
from body import *

today =  date.today()
day = today.strftime("%d")
month = today.strftime("%m")
monthNa = today.strftime("%B")
monthAb = today.strftime("%b")
year = today.strftime("%Y")
yearSh = today.strftime("%y")

timeTup = (day,month,monthNa,monthAb,year,yearSh)

def add_hyperlink(paragraph, text, url):
  #https://github.com/python-openxml/python-docx/issues/384
  # This gets access to the document.xml.rels file and gets a new relation id value
  part = paragraph.part
  r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

  # Create the w:hyperlink tag and add needed values
  hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
  hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

  # Create a w:r element and a new w:rPr element
  new_run = docx.oxml.shared.OxmlElement('w:r')
  rPr = docx.oxml.shared.OxmlElement('w:rPr')

  # Join all the xml elements together add add the required text to the w:r element
  new_run.append(rPr)
  new_run.text = text
  hyperlink.append(new_run)

  # Create a new Run object and add the hyperlink into it
  r = paragraph.add_run ()
  r._r.append (hyperlink)

  # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
  # Delete this if using a template that has the hyperlink style in it
  r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
  r.font.underline = True

  return hyperlink

def makeEndStub(doc,sizeKey):

  doc.add_page_break()

  cleanItemData = OrderedDict()
  cleanLines = []
  with open(sizeKey, "rt") as f:
    for lines in f:
      cleanLines.append(lines)
  for lines in cleanLines:
    #Inseam|31"|legs|1|
    line = lines.split("|")
    lineData = OrderedDict()
    lineData["Type"] = line[2]
    lineData["Sub"] = line[3]
    lineData["Value"] = line[1]
    cleanItemData[line[0]] = lineData
  cleanItemData = OrderedDict(sorted(cleanItemData.items(), key=lambda t: t[0]))
  cleanItemData = OrderedDict(sorted(cleanItemData.items(), key=lambda t: t[1]["Sub"][0]))
  cleanItemData = OrderedDict(sorted(cleanItemData.items(), key=lambda t: t[1]["Type"][0]))
  table = doc.add_table(rows=1, cols=2)
  for key in cleanItemData:
    cells = table.add_row().cells
    if(cleanItemData[key]["Type"][2:]!="space"):
      cells[0].text = key
      cells[1].text = cleanItemData[key]["Value"]


def makeDoc(itemDict,isSized=False):
  doc = Document()
  styles = doc.styles

  sec = doc.sections[0]
  h1 = sec.header.paragraphs[0]
  h1.text = ((titleDict["header"]).format(*timeTup))
  h1.style = styles["Heading 2"]
  h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

  lastType = "0 Appearal"
  lastKey = ""
  type = lastType
  i = 0
  t = 0
  imgAdr0 = ""
  imgAdr1 = ""
  itmAdr0 = ""
  itmAdr1 = ""

  title = doc.add_paragraph(type[2:])
  titformat = title.paragraph_format
  titformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
  titformat.tab_stops.add_tab_stop(Inches(3))

  for key in itemDict:
    type = itemDict[key]["Type"]
    if(type==lastType):
      if(i==0):
        desc = []
        desc.append("{}".format(itemDict[key]["Name"]))
        desc.append("{0:.2f}{1}".format(itemDict[key]["Price"],getShipping(itemDict[key])))
        imgAdr0 = itemDict[key]["Image"]
        itmAdr0 = itemDict[key]["Link"]
        i = 1
      else:
        desc.append("{0}".format(itemDict[key]["Name"]))
        desc.append("{0:.2f}{1}".format(itemDict[key]["Price"],getShipping(itemDict[key])))
        imgAdr1 = itemDict[key]["Image"]
        itmAdr1 = itemDict[key]["Link"]
        i = 0
        para = doc.add_paragraph()
        paraformat = para.paragraph_format
        paraformat.tab_stops.add_tab_stop(Inches(3))
        paraformat.keep_together = True
        add_hyperlink(para,desc[0],itmAdr0)
        para.add_run("\t")
        add_hyperlink(para,desc[2],itmAdr1)
        run = para.add_run()
        run.add_break()
        run.add_picture(imgAdr0, height=Inches(1))
        run.add_tab()
        run.add_picture(imgAdr1,height=Inches(1))
        run.add_break()
        run.add_text(desc[1])
        run.add_tab()
        run.add_text(desc[3])
        run.add_break()
    else:
      if(i==1):
        i = 0
        para = doc.add_paragraph()
        paraformat = para.paragraph_format
        paraformat.tab_stops.add_tab_stop(Inches(3))
        paraformat.keep_together = True
        add_hyperlink(para,desc[0],itmAdr0)
        run = para.add_run()
        run.add_break()
        run.add_picture(imgAdr0, height=Inches(1))
        run.add_break()
        run.add_text(desc[1])
        title = doc.add_paragraph(type[2:]+"\t")
        titformat = title.paragraph_format
        titformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titformat.tab_stops.add_tab_stop(Inches(3))
        desc = []
        desc.append("{}".format(itemDict[key]["Name"]))
        desc.append("{0:.2f}{1}".format(itemDict[key]["Price"],getShipping(itemDict[key])))
        imgAdr0 = itemDict[key]["Image"]
        i = 1
      else:
        title = doc.add_paragraph(type[2:]+"\t")
        titformat = title.paragraph_format
        titformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titformat.tab_stops.add_tab_stop(Inches(3))
        desc = []
        desc.append("{}".format(itemDict[key]["Name"]))
        desc.append("{0:.2f}{1}".format(itemDict[key]["Price"],getShipping(itemDict[key])))
        imgAdr0 = itemDict[key]["Image"]
        i = 1
    lastType = type
    lastKey = key
    t+=1

  if(t%2):
    desc = []
    desc.append("{}".format(itemDict[key]["Name"]))
    desc.append("{0:.2f}{1}".format(itemDict[key]["Price"],getShipping(itemDict[key])))
    imgAdr0 = itemDict[key]["Image"]
    itmAdr0 = itemDict[key]["Link"]
    para = doc.add_paragraph()
    paraformat = para.paragraph_format
    paraformat.tab_stops.add_tab_stop(Inches(3))
    paraformat.keep_together = True
    add_hyperlink(para,desc[0],itmAdr0)
    run = para.add_run()
    run.add_break()
    run.add_picture(imgAdr0, height=Inches(1))
    run.add_break()
    run.add_text(desc[1])

  if(isSized):
    sizeKey = "./data/sizeKey.txt"
    makeEndStub(doc,sizeKey)

  doc.save("{}.docx".format(titleDict["file"]))


def makeText(itemDict,isSized=False):
  doc = Document()
  styles = doc.styles

  sec = doc.sections[0]
  h1 = sec.header.paragraphs[0]
  h1.text = ((titleDict["header"]).format(*timeTup))
  h1.style = styles["Heading 2"]
  h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

  lastType = "0 Appearal"
  type = lastType
  itmAdr0 = ""
  itmAdr1 = ""

  title = doc.add_paragraph(type[2:])
  titformat = title.paragraph_format
  titformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
  titformat.tab_stops.add_tab_stop(Inches(3))

  for key in itemDict:
    type = itemDict[key]["Type"]
    if(type!=lastType):
      title = doc.add_paragraph(type[2:])
      titformat = title.paragraph_format
      titformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph("{} –\t".format(itemDict[key]["Name"]))
    paraformat = para.paragraph_format
    paraformat.tab_stops.add_tab_stop(Inches(3))
    paraformat.tab_stops.add_tab_stop(Inches(4))
    paraformat.keep_together = True
    add_hyperlink(para,"Link",itemDict[key]["Link"])
    run = para.add_run()
    run.add_text("\t– {0:.2f}{1}".format(itemDict[key]["Price"],getShipping(itemDict[key])))
    run.add_break()
    lastType = type

  if(isSized):
    sizeKey = "./data/sizeKey.txt"
    makeEndStub(doc,sizeKey)

  doc.save("{}_text.docx".format(titleDict["file"]))

makeDoc(itemDict,True)
